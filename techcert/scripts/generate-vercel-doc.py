#!/usr/bin/env python3
"""Generate Word documentation for Vercel frontend-backend setup."""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

OUTPUT = Path("/home/samuel/Desktop/Vercel-Frontend-Backend-Guide.docx")


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = cell
    doc.add_paragraph()


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Connecting Frontend to Backend on Vercel", 0)
    doc.add_paragraph(
        "Complete deployment guide for SignalForge / DoraHacks: Next.js frontend + "
        "Express backend on Vercel, with MongoDB Atlas. Updated with real project URLs "
        "and troubleshooting steps from production deployment."
    )

    doc.add_heading("Architecture Overview", 1)
    add_code(
        doc,
        "GitHub: samuelkwartengokyere/Dora-Hacks\n"
        "├── techcert/frontend/  →  Vercel: dora-hacks (frontend)\n"
        "└── techcert/backend/   →  Vercel: dora-hacks-j399 (backend)",
    )
    add_code(
        doc,
        "User browser\n"
        "    ↓\n"
        "Frontend (dora-hacks.vercel.app)\n"
        "    ↓  NEXT_PUBLIC_API_URL\n"
        "Backend (dora-hacks-j399.vercel.app/api)\n"
        "    ↓  MONGODB_URI\n"
        "MongoDB Atlas (cloud database)",
    )

    doc.add_heading("Critical Rule: What Goes Where", 1)
    add_table(
        doc,
        ["Variable", "Frontend (dora-hacks)", "Backend (dora-hacks-j399)"],
        [
            ["NEXT_PUBLIC_API_URL", "YES — backend HTTP URL", "NO"],
            ["MONGODB_URI", "NO — never put this on frontend", "YES — Atlas connection string"],
            ["JWT_SECRET", "NO", "YES"],
            ["ADMIN_EMAIL / ADMIN_PASSWORD", "NO", "YES (recommended)"],
            ["FRONTEND_URL", "NO", "Optional"],
        ],
    )
    doc.add_paragraph(
        "Common mistake: putting the MongoDB connection string in NEXT_PUBLIC_API_URL on the "
        "frontend. The frontend must use an https://...vercel.app/api URL, not mongodb+srv://..."
    )

    doc.add_heading("Project URLs (Your Deployment)", 1)
    add_table(
        doc,
        ["Project", "Vercel Name", "URL"],
        [
            ["Backend", "dora-hacks-j399", "https://dora-hacks-j399.vercel.app"],
            ["Frontend", "dora-hacks", "https://dora-hacks.vercel.app"],
        ],
    )

    doc.add_heading("Part 1 — Backend Setup (dora-hacks-j399)", 1)

    doc.add_heading("1.1 Connect GitHub", 2)
    for step in [
        "Open dora-hacks-j399 on Vercel",
        "Settings → Git → Connect to samuelkwartengokyere/Dora-Hacks",
        "Production branch: main",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("1.2 Root Directory", 2)
    add_code(doc, "techcert/backend")
    add_table(
        doc,
        ["Setting", "Value"],
        [
            ["Root Directory", "techcert/backend"],
            ["Framework", "Other / Node.js"],
            ["Build Command", "(empty)"],
            ["Output Directory", "(empty)"],
            ["Install Command", "npm install"],
        ],
    )

    doc.add_heading("1.3 MongoDB Atlas Setup", 2)
    for step in [
        "Create free M0 cluster at mongodb.com/atlas",
        "Database Access → Add user (e.g. signalforge) → Autogenerate password → Copy",
        "Network Access → Add IP Address → Allow Access from Anywhere (0.0.0.0/0)",
        "Database → Connect → Drivers → Node.js → copy connection string",
        "Replace <db_password> with your database user password",
        "Add /signalforge before the ? in the connection string",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_paragraph("Correct connection string format:")
    add_code(
        doc,
        "mongodb+srv://signalforge:YOUR_PASSWORD@cluster0.iiiu7hd.mongodb.net/"
        "signalforge?retryWrites=true&w=majority",
    )

    doc.add_heading("1.4 Backend Environment Variables", 2)
    doc.add_paragraph("Settings → Environment Variables on dora-hacks-j399 ONLY:")
    add_table(
        doc,
        ["Key", "Value"],
        [
            ["MONGODB_URI", "Full Atlas connection string (see above)"],
            ["JWT_SECRET", "Long random string (32+ characters)"],
            ["ADMIN_EMAIL", "admin@signalforge.ai"],
            ["ADMIN_PASSWORD", "admin123"],
            ["FRONTEND_URL", "https://dora-hacks.vercel.app (optional)"],
        ],
    )
    doc.add_paragraph("Rules for MONGODB_URI:")
    for rule in [
        "No quotes around the value",
        "One line, no spaces at start or end",
        "Must include /signalforge before the ?",
        "Use the database user password, not your Atlas account login",
        "Enable Production, Preview, and Development",
    ]:
        doc.add_paragraph(rule, style="List Bullet")

    doc.add_heading("1.5 Redeploy Backend", 2)
    for step in [
        "Deployments → ⋯ → Redeploy",
        "Uncheck Use existing Build Cache",
        "Wait until status is Ready",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("1.6 Verify Backend", 2)
    doc.add_paragraph("Basic health (no database required):")
    add_code(doc, "https://dora-hacks-j399.vercel.app/api/health")
    doc.add_paragraph('Expected: { "status": "ok", "service": "SignalForge AI API", ... }')
    doc.add_paragraph("Database diagnostic (use this to verify MONGODB_URI):")
    add_code(doc, "https://dora-hacks-j399.vercel.app/api/health/db")
    doc.add_paragraph("Success response:")
    add_code(
        doc,
        '{\n  "connected": true,\n  "config": {\n'
        '    "user": "signalforge",\n'
        '    "host": "cluster0.iiiu7hd.mongodb.net",\n'
        '    "database": "signalforge",\n'
        '    "passwordLength": 20\n  },\n'
        '  "message": "MongoDB connection successful."\n}',
    )
    doc.add_paragraph("If connected is false, compare config fields:")
    add_table(
        doc,
        ["Field", "Wrong (broken)", "Correct (working)"],
        [
            ["user", "samfine278_db_user (old user)", "signalforge"],
            ["database", "null", "signalforge"],
            ["passwordLength", "13 (truncated/wrong password)", "matches Atlas password length"],
            ["uriLength", "~93 (too short)", "~120+ (full string)"],
        ],
    )

    doc.add_heading("Part 2 — Frontend Setup (dora-hacks)", 1)

    doc.add_heading("2.1 Connect GitHub & Root Directory", 2)
    add_code(doc, "Root Directory: techcert/frontend")

    doc.add_heading("2.2 Frontend Environment Variables", 2)
    doc.add_paragraph("Settings → Environment Variables on dora-hacks ONLY:")
    add_table(
        doc,
        ["Key", "Value", "Notes"],
        [
            ["NEXT_PUBLIC_API_URL", "https://dora-hacks-j399.vercel.app/api", "Required"],
            ["MONGODB_URI", "DELETE if present", "Never on frontend"],
        ],
    )

    doc.add_heading("2.3 Redeploy Frontend", 2)
    doc.add_paragraph(
        "NEXT_PUBLIC_* variables are baked in at build time. You MUST redeploy after "
        "adding or changing NEXT_PUBLIC_API_URL."
    )
    for step in [
        "Deployments → ⋯ → Redeploy",
        "Uncheck Use existing Build Cache",
        "Wait until Ready",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Part 3 — Test Login", 1)
    doc.add_paragraph("1. Confirm /api/health/db shows connected: true on backend")
    doc.add_paragraph("2. Open https://dora-hacks.vercel.app/admin")
    doc.add_paragraph("3. Sign in with:")
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Email", "admin@signalforge.ai"],
            ["Password", "admin123"],
        ],
    )
    doc.add_paragraph("Or use Create account to register a new user.")

    doc.add_heading("Part 4 — Verify Frontend → Backend Connection", 1)
    doc.add_paragraph("Browser DevTools → Network tab → try login → check Request URL:")
    add_table(
        doc,
        ["Request URL", "Meaning"],
        [
            ["https://dora-hacks-j399.vercel.app/api/auth/login", "Correct"],
            ["http://localhost:5000/api/auth/login", "NEXT_PUBLIC_API_URL not set or not redeployed"],
            ["mongodb+srv://...", "MongoDB string wrongly set as NEXT_PUBLIC_API_URL"],
        ],
    )

    doc.add_heading("Part 5 — Troubleshooting", 1)
    add_table(
        doc,
        ["Error", "Cause", "Fix"],
        [
            ["Failed to fetch", "Frontend cannot reach API", "Set NEXT_PUBLIC_API_URL; redeploy frontend"],
            ["Cannot reach API at mongodb+srv://...", "MongoDB URI on frontend", "Delete MONGODB_URI from frontend; set NEXT_PUBLIC_API_URL to backend URL"],
            ["Cannot reach API at localhost", "Env var missing", "Set NEXT_PUBLIC_API_URL; redeploy frontend"],
            ["MONGODB_URI is not set", "Backend missing DB", "Add MONGODB_URI on backend; redeploy"],
            ["IP not whitelisted", "Atlas blocking Vercel", "Atlas → Network Access → 0.0.0.0/0"],
            ["bad auth : authentication failed", "Wrong password or user on Vercel", "Delete MONGODB_URI on backend; re-add with signalforge user; redeploy; check /api/health/db"],
            ["user still samfine278_db_user in /api/health/db", "Old env var still deployed", "Delete MONGODB_URI, add new value, redeploy without cache"],
            ["database: null in /api/health/db", "Missing /signalforge in URI", "Add /signalforge before ? in connection string"],
            ["passwordLength too short", "Truncated or wrong password", "Re-copy full password from Atlas; no quotes"],
            ["Invalid credentials", "API works; wrong login", "Use admin@signalforge.ai / admin123 or Register"],
            ["500 FUNCTION_INVOCATION_FAILED", "Backend crash", "Check Deployments → Logs on backend"],
        ],
    )

    doc.add_heading("Part 6 — Fix bad auth (Step by Step)", 1)
    for step in [
        "Atlas → Database Access → user signalforge → Edit Password → Autogenerate → Copy",
        "Build: mongodb+srv://signalforge:PASSWORD@cluster0.iiiu7hd.mongodb.net/signalforge?retryWrites=true&w=majority",
        "Vercel dora-hacks-j399 → Environment Variables → DELETE all MONGODB_URI entries",
        "Add new MONGODB_URI (no quotes, all environments checked) → Save",
        "Deployments → Redeploy (uncheck build cache) → Wait for Ready",
        "Open /api/health/db — must show connected: true before trying login",
        "If frontend had MONGODB_URI, delete it and redeploy frontend too",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Part 7 — Auto-Deploy on Git Push", 1)
    doc.add_paragraph(
        "With Git connected, pushing to main triggers new deployments. Environment variable "
        "changes require a manual Redeploy to take effect."
    )

    doc.add_heading("Part 8 — Local Development", 1)
    doc.add_paragraph("frontend/.env.local:")
    add_code(doc, "NEXT_PUBLIC_API_URL=http://localhost:5000/api")
    doc.add_paragraph("backend/.env:")
    add_code(
        doc,
        "MONGODB_URI=mongodb://localhost:27017/signalforge\n"
        "JWT_SECRET=local-dev-secret\n"
        "ADMIN_EMAIL=admin@signalforge.ai\n"
        "ADMIN_PASSWORD=admin123",
    )
    doc.add_paragraph("Run: npm run dev (from techcert folder)")

    doc.add_heading("Part 9 — Complete Checklist", 1)

    doc.add_heading("Backend (dora-hacks-j399)", 2)
    for item in [
        "Git connected to samuelkwartengokyere/Dora-Hacks",
        "Root Directory = techcert/backend",
        "MONGODB_URI with user signalforge and /signalforge database",
        "JWT_SECRET, ADMIN_EMAIL, ADMIN_PASSWORD set",
        "Atlas Network Access includes 0.0.0.0/0",
        "/api/health returns ok",
        "/api/health/db returns connected: true",
        "Backend redeployed after env changes",
    ]:
        doc.add_paragraph(f"☐ {item}")

    doc.add_heading("Frontend (dora-hacks)", 2)
    for item in [
        "Git connected to same repo",
        "Root Directory = techcert/frontend",
        "NEXT_PUBLIC_API_URL = https://dora-hacks-j399.vercel.app/api",
        "MONGODB_URI deleted if it was mistakenly added",
        "Frontend redeployed after env changes",
        "Login request goes to backend URL in Network tab",
    ]:
        doc.add_paragraph(f"☐ {item}")

    doc.add_heading("Quick Copy-Paste", 1)
    doc.add_paragraph("Frontend (dora-hacks):")
    add_code(doc, "NEXT_PUBLIC_API_URL=https://dora-hacks-j399.vercel.app/api")
    doc.add_paragraph("Backend (dora-hacks-j399):")
    add_code(
        doc,
        "MONGODB_URI=mongodb+srv://signalforge:PASSWORD@cluster0.iiiu7hd.mongodb.net/"
        "signalforge?retryWrites=true&w=majority\n"
        "JWT_SECRET=your-long-random-secret\n"
        "ADMIN_EMAIL=admin@signalforge.ai\n"
        "ADMIN_PASSWORD=admin123",
    )
    doc.add_paragraph("After any env change: redeploy that project (uncheck build cache).")

    doc.add_paragraph()
    doc.add_paragraph("SignalForge AI / DoraHacks — Vercel Deployment Guide (updated).")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(f"Created: {path}")
